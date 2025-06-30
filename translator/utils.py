import os
import pytesseract
from PIL import Image
import deepl
from pdf2image import convert_from_path
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from django.conf import settings
from io import BytesIO
import tempfile

def perform_ocr(file_path, language='eng'):
    """
    Perform OCR on the given file.
    
    Args:
        file_path (str): Path to the file
        language (str): Language code for OCR
        
    Returns:
        str: Extracted text
    """
    # Set Tesseract command path from settings
    pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # For PDF files, convert to images first
    if file_ext == '.pdf':
        return ocr_pdf(file_path, language)
    
    # For image files
    elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
        return ocr_image(file_path, language)
    
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")

def ocr_image(image_path, language='eng'):
    """
    Perform OCR on an image file.
    
    Args:
        image_path (str): Path to the image file
        language (str): Language code for OCR
        
    Returns:
        str: Extracted text
    """
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang=language)
        return text
    except Exception as e:
        raise Exception(f"Error performing OCR on image: {str(e)}")

def ocr_pdf(pdf_path, language='eng'):
    """
    Perform OCR on a PDF file by converting it to images.
    
    Args:
        pdf_path (str): Path to the PDF file
        language (str): Language code for OCR
        
    Returns:
        str: Extracted text
    """
    try:
        # Convert PDF to images
        images = convert_from_path(pdf_path)
        
        # Perform OCR on each image
        text = ""
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img, lang=language)
            text += f"\n\n--- Page {i+1} ---\n\n{page_text}"
        
        return text
    except Exception as e:
        raise Exception(f"Error performing OCR on PDF: {str(e)}")

def translate_text(text, source_lang=None, target_lang='EN-US'):
    """
    Translate text using DeepL API.
    
    Args:
        text (str): Text to translate
        source_lang (str, optional): Source language code
        target_lang (str): Target language code
        
    Returns:
        str: Translated text
    """
    try:
        # Initialize DeepL translator
        translator = deepl.Translator(settings.DEEPL_API_KEY)
        
        # Translate text
        result = translator.translate_text(
            text,
            source_lang=source_lang,
            target_lang=target_lang
        )
        
        return result.text
    except Exception as e:
        raise Exception(f"Error translating text: {str(e)}")

def create_text_file(text, output_path):
    """
    Create a text file with the given text.
    
    Args:
        text (str): Text content
        output_path (str): Path to save the file
        
    Returns:
        str: Path to the created file
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return output_path
    except Exception as e:
        raise Exception(f"Error creating text file: {str(e)}")

def create_docx_file(text, output_path):
    """
    Create a Word document with the given text.
    
    Args:
        text (str): Text content
        output_path (str): Path to save the file
        
    Returns:
        str: Path to the created file
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        doc = Document()
        
        # Set font properties for better Japanese support
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial Unicode MS'  # This font supports Japanese
        font.size = Pt(11)
        
        # Split text by page markers
        pages = text.split("--- Page")
        
        # Add first part (if any)
        if pages[0].strip():
            paragraph = doc.add_paragraph(pages[0].strip())
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Add remaining pages with page breaks
        for page in pages[1:]:
            # Extract page number and content
            parts = page.split("---", 1)
            if len(parts) > 1:
                content = parts[1].strip()
            else:
                content = parts[0].strip()
            
            # Add page break and content
            doc.add_page_break()
            paragraph = doc.add_paragraph(content)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        doc.save(output_path)
        return output_path
    except Exception as e:
        # Try a simpler approach if the first one fails
        try:
            from docx import Document
            
            doc = Document()
            
            # Split text by page markers
            pages = text.split("--- Page")
            
            # Add first part (if any)
            if pages[0].strip():
                doc.add_paragraph(pages[0].strip())
            
            # Add remaining pages with page breaks
            for page in pages[1:]:
                # Extract page number and content
                parts = page.split("---", 1)
                if len(parts) > 1:
                    content = parts[1].strip()
                else:
                    content = parts[0].strip()
                
                # Add page break and content
                doc.add_page_break()
                doc.add_paragraph(content)
            
            doc.save(output_path)
            return output_path
            
        except Exception as nested_e:
            raise Exception(f"Error creating Word document: {str(e)} -> {str(nested_e)}")

def create_pdf_file(text, output_path):
    """
    Create a PDF file with the given text.
    
    Args:
        text (str): Text content
        output_path (str): Path to save the file
        
    Returns:
        str: Path to the created file
    """
    try:
        # Use FPDF2 which has better Unicode support
        from fpdf import FPDF
        
        # Create PDF with Unicode support
        pdf = FPDF()
        pdf.add_page()
        
        # Add a Unicode font that supports Japanese
        # First try to use an installed font that supports Japanese
        try:
            # Try to use Arial Unicode MS if available
            pdf.add_font('Arial Unicode MS', '', '', uni=True)
            pdf.set_font('Arial Unicode MS', '', 10)
        except RuntimeError:
            try:
                # Try to use MS Gothic if available (common on Windows)
                pdf.add_font('MS Gothic', '', '', uni=True)
                pdf.set_font('MS Gothic', '', 10)
            except RuntimeError:
                try:
                    # Try to use Noto Sans CJK if available
                    pdf.add_font('Noto Sans CJK JP', '', '', uni=True)
                    pdf.set_font('Noto Sans CJK JP', '', 10)
                except RuntimeError:
                    # Fall back to the default font with limited Unicode support
                    pdf.set_font('Arial', '', 10)
        
        # Split text by page markers
        pages = text.split("--- Page")
        
        # Process each page
        for i, page_content in enumerate(pages):
            if i > 0:  # Not the first item
                pdf.add_page()  # Start a new page
            
            # Clean up the page content
            if i == 0:  # First part
                content = page_content.strip()
            else:
                # Extract content after the page number
                parts = page_content.split("---", 1)
                if len(parts) > 1:
                    content = parts[1].strip()
                else:
                    content = parts[0].strip()
            
            # Split content into lines and write to PDF
            lines = content.split('\n')
            for line in lines:
                # Handle empty lines
                if not line.strip():
                    pdf.ln(5)  # Add some space
                    continue
                
                # Write the line with multi-cell to handle wrapping
                pdf.multi_cell(0, 5, line)
                pdf.ln(2)  # Add a small space between lines
        
        # Save the PDF
        pdf.output(output_path)
        return output_path
    except Exception as e:
        # Try alternative method if the first one fails
        try:
            # Use ReportLab with a Unicode font
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import tempfile
            import os
            
            # Try to register a font that supports Japanese
            font_registered = False
            
            # List of potential fonts that support Japanese
            potential_fonts = [
                ('Arial Unicode MS', 'arialuni.ttf'),
                ('MS Gothic', 'msgothic.ttc'),
                ('Yu Gothic', 'yugothic.ttf'),
                ('Meiryo', 'meiryo.ttc'),
                ('Noto Sans CJK JP', 'NotoSansCJKjp-Regular.otf')
            ]
            
            # Common font locations on Windows
            font_dirs = [
                'C:/Windows/Fonts/',
                os.path.expanduser('~/.fonts/'),
                os.path.expanduser('~/Library/Fonts/')
            ]
            
            # Try to find and register a suitable font
            for font_name, font_file in potential_fonts:
                for font_dir in font_dirs:
                    font_path = os.path.join(font_dir, font_file)
                    if os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont(font_name, font_path))
                            font_registered = True
                            break
                        except:
                            continue
                if font_registered:
                    break
            
            # Create a PDF document
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # Set font - use registered font or default
            if font_registered:
                c.setFont(font_name, 10)
            else:
                c.setFont('Helvetica', 10)
            
            # Split text by page markers
            pages = text.split("--- Page")
            
            # Process each page
            for i, page_content in enumerate(pages):
                if i > 0:  # Not the first item
                    c.showPage()  # Start a new page
                    if font_registered:
                        c.setFont(font_name, 10)
                    else:
                        c.setFont('Helvetica', 10)
                
                # Clean up the page content
                if i == 0:  # First part
                    content = page_content.strip()
                else:
                    # Extract content after the page number
                    parts = page_content.split("---", 1)
                    if len(parts) > 1:
                        content = parts[1].strip()
                    else:
                        content = parts[0].strip()
                
                # Split content into lines
                lines = content.split('\n')
                y = height - 40  # Start from top with margin
                
                for line in lines:
                    if y < 40:  # Bottom margin
                        c.showPage()
                        if font_registered:
                            c.setFont(font_name, 10)
                        else:
                            c.setFont('Helvetica', 10)
                        y = height - 40
                    
                    # Draw the text
                    c.drawString(40, y, line)
                    y -= 14  # Line spacing
            
            c.save()
            return output_path
            
        except Exception as nested_e:
            # If both methods fail, create a simple text file and convert it to PDF
            try:
                import subprocess
                
                # Create a temporary text file
                with tempfile.NamedTemporaryFile(suffix='.txt', delete=False, mode='w', encoding='utf-8') as temp_file:
                    temp_file.write(text)
                    temp_txt_path = temp_file.name
                
                # Try to use an external tool to convert text to PDF
                # This is a fallback method and might not work on all systems
                try:
                    # Try using wkhtmltopdf if available
                    html_content = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <style>
                            body {{ font-family: Arial, sans-serif; }}
                            pre {{ white-space: pre-wrap; }}
                        </style>
                    </head>
                    <body>
                        <pre>{text}</pre>
                    </body>
                    </html>
                    """
                    
                    # Create a temporary HTML file
                    with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8') as html_file:
                        html_file.write(html_content)
                        temp_html_path = html_file.name
                    
                    # Try to convert HTML to PDF using wkhtmltopdf
                    subprocess.run(['wkhtmltopdf', temp_html_path, output_path], check=True)
                    
                    # Clean up temporary files
                    os.unlink(temp_html_path)
                    os.unlink(temp_txt_path)
                    
                    return output_path
                    
                except (subprocess.SubprocessError, FileNotFoundError):
                    # If wkhtmltopdf is not available, just return the text file
                    import shutil
                    shutil.copy(temp_txt_path, output_path.replace('.pdf', '.txt'))
                    os.unlink(temp_txt_path)
                    return output_path.replace('.pdf', '.txt')
                    
            except Exception as final_e:
                raise Exception(f"Error creating PDF file: {str(e)} -> {str(nested_e)} -> {str(final_e)}")
    
        raise Exception(f"Error creating PDF file: {str(e)}")

def create_image_file(text, output_path, format='JPEG'):
    """
    Create an image with the given text.
    
    Args:
        text (str): Text content
        output_path (str): Path to save the file
        format (str): Image format (JPEG or PNG)
        
    Returns:
        str: Path to the created file
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
        import os
        
        # Create a white background image
        width, height = 1000, 1400
        img = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)
        
        # Try to find a font that supports Japanese characters
        font = None
        
        # List of potential fonts that support Japanese
        potential_fonts = [
            ('Arial Unicode MS', 'arialuni.ttf'),
            ('MS Gothic', 'msgothic.ttc'),
            ('Yu Gothic', 'yugothic.ttf'),
            ('Meiryo', 'meiryo.ttc'),
            ('Noto Sans CJK JP', 'NotoSansCJKjp-Regular.otf')
        ]
        
        # Common font locations on Windows
        font_dirs = [
            'C:/Windows/Fonts/',
            os.path.expanduser('~/.fonts/'),
            os.path.expanduser('~/Library/Fonts/')
        ]
        
        # Try to find a suitable font
        for font_name, font_file in potential_fonts:
            for font_dir in font_dirs:
                font_path = os.path.join(font_dir, font_file)
                if os.path.exists(font_path):
                    try:
                        font = ImageFont.truetype(font_path, 14)
                        break
                    except:
                        continue
            if font:
                break
        
        # If no suitable font found, use default
        if not font:
            try:
                font = ImageFont.truetype("arial.ttf", 14)
            except:
                font = ImageFont.load_default()
        
        # Draw text on the image
        lines = text.split('\n')
        y = 20  # Start from top with margin
        
        for line in lines:
            # Handle long lines by wrapping text
            if font.getlength(line) > width - 40:
                # Simple word wrapping
                words = line.split()
                current_line = ""
                
                for word in words:
                    test_line = current_line + " " + word if current_line else word
                    if font.getlength(test_line) <= width - 40:
                        current_line = test_line
                    else:
                        # Draw current line and start a new one
                        draw.text((20, y), current_line, fill='black', font=font)
                        y += 20
                        current_line = word
                        
                        # Check if we need a new page
                        if y > height - 20:
                            # Save current image
                            img.save(output_path, format=format)
                            
                            # Create a new image for the rest of the text
                            img = Image.new('RGB', (width, height), color='white')
                            draw = ImageDraw.Draw(img)
                            y = 20
                
                # Draw the last line if there's anything left
                if current_line:
                    draw.text((20, y), current_line, fill='black', font=font)
                    y += 20
            else:
                # Draw the line directly
                draw.text((20, y), line, fill='black', font=font)
                y += 20
            
            # Check if we need a new page
            if y > height - 20:
                # Save current image
                img.save(output_path, format=format)
                
                # Create a new image for the rest of the text
                img = Image.new('RGB', (width, height), color='white')
                draw = ImageDraw.Draw(img)
                y = 20
        
        # Save the image
        img.save(output_path, format=format)
        return output_path
    except Exception as e:
        raise Exception(f"Error creating image file: {str(e)}")

def generate_output_file(text, output_format, output_path):
    """
    Generate an output file in the specified format.
    
    Args:
        text (str): Text content
        output_format (str): Output format (pdf, docx, txt, jpeg, png)
        output_path (str): Path to save the file
        
    Returns:
        str: Path to the created file
    """
    if output_format == 'txt':
        return create_text_file(text, output_path)
    elif output_format == 'docx':
        return create_docx_file(text, output_path)
    elif output_format == 'pdf':
        return create_pdf_file(text, output_path)
    elif output_format == 'jpeg':
        return create_image_file(text, output_path, format='JPEG')
    elif output_format == 'png':
        return create_image_file(text, output_path, format='PNG')
    else:
        raise ValueError(f"Unsupported output format: {output_format}")